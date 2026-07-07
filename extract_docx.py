import docx
import sys

def extract_text(doc_path, txt_path):
    try:
        doc = docx.Document(doc_path)
        with open(txt_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.replace('\n', ' ') for cell in row.cells]
                    f.write('\t'.join(row_data) + '\n')
        print(f"Extracted to {txt_path}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    extract_text(sys.argv[1], sys.argv[2])
